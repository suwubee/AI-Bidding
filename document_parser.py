import os
import re
from typing import Dict, List, Any
import PyPDF2
from docx import Document
from docx.shared import Pt
import chardet

class DocumentParser:
    """文档解析器，支持PDF、DOC、DOCX格式"""
    
    def __init__(self):
        self.heading_patterns = [
            # 级别1：章节
            r'^第[一二三四五六七八九十\d]+章[\s\u3000]*[^\d\s].*',
            # 级别2：节
            r'^第[一二三四五六七八九十\d]+节[\s\u3000]*[^\d\s].*',
            # 级别3：条款
            r'^第[一二三四五六七八九十\d]+条[\s\u3000]*[^\d\s].*',
            # 级别4：数字编号（如：1. 标题）
            r'^\d+\.[\s\u3000]*[^\d\s].*',
            # 级别5：二级数字编号（如：1.1 标题）
            r'^\d+\.\d+[\s\u3000]*[^\d\s].*',
            # 级别6：三级数字编号（如：1.1.1 标题）
            r'^\d+\.\d+\.\d+[\s\u3000]*[^\d\s].*',
            # 级别7：四级数字编号（如：1.1.1.1 标题）
            r'^\d+\.\d+\.\d+\.\d+[\s\u3000]*[^\d\s].*',
        ]
        
        # 额外的标题模式（优先级较低）
        self.secondary_patterns = [
            r'^[一二三四五六七八九十]+、[\s\u3000]*[^\s].*',
            r'^（[一二三四五六七八九十]+）[\s\u3000]*[^\s].*',
            r'^\([一二三四五六七八九十]+\)[\s\u3000]*[^\s].*',
            r'^[A-Z]\.[\s\u3000]*[^\s].*',
            r'^\([A-Z]\)[\s\u3000]*[^\s].*'
        ]
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """解析文档并返回结构化信息"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._parse_pdf(file_path)
        elif file_extension in ['.doc', '.docx']:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}")
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文档"""
        result = {
            'document_type': 'PDF',
            'total_pages': 0,
            'headings': [],
            'structure': [],
            'content_preview': ''
        }
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                result['total_pages'] = len(pdf_reader.pages)
                
                print(f"\n=== 开始解析PDF文档，共{result['total_pages']}页 ===")
                
                # 策略1: 优先尝试书签解析（PDF的最佳方案）
                headings_from_bookmarks = self._extract_pdf_bookmarks_intelligent(pdf_reader)
                
                # 策略2: 如果书签无效，提取文本内容进行分析
                full_text = ""
                if not headings_from_bookmarks:
                    print("书签解析未找到有效结构，提取文本内容...")
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            text = page.extract_text()
                            full_text += text + "\n"
                        except:
                            continue
                    
                    # 使用文本分析
                    headings_from_text = self._extract_headings_from_text(full_text)
                else:
                    # 即使有书签，也提取一些文本用于预览
                    try:
                        for page_num, page in enumerate(pdf_reader.pages[:3]):  # 只提取前3页用于预览
                            text = page.extract_text()
                            full_text += text + "\n"
                            if len(full_text) > 1000:  # 够用就行
                                break
                    except:
                        pass
                
                # 选择最佳解析结果
                if headings_from_bookmarks:
                    result['headings'] = headings_from_bookmarks
                    result['extraction_method'] = 'bookmarks'
                    print(f"使用书签解析，找到{len(headings_from_bookmarks)}个标题")
                elif 'headings_from_text' in locals() and headings_from_text:
                    result['headings'] = headings_from_text
                    result['extraction_method'] = 'text_analysis'
                    print(f"使用文本分析，找到{len(headings_from_text)}个标题")
                else:
                    result['headings'] = []
                    result['extraction_method'] = 'no_structure_found'
                    print("未找到任何文档结构")
                
                print(f"=== PDF解析完成，最终找到{len(result['headings'])}个标题 ===\n")
                
                result['structure'] = self._build_document_structure(result['headings'])
                result['content_preview'] = full_text[:500] + "..." if len(full_text) > 500 else full_text
                
        except Exception as e:
            raise Exception(f"PDF解析错误: {str(e)}")
        
        return result
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析DOCX文档 - 智能多策略解析"""
        result = {
            'document_type': 'DOCX',
            'total_paragraphs': 0,
            'headings': [],
            'structure': [],
            'content_preview': ''
        }
        
        try:
            doc = Document(file_path)
            result['total_paragraphs'] = len(doc.paragraphs)
            
            print(f"\n=== 开始解析DOCX文档，共{len(doc.paragraphs)}个段落 ===")
            
            # 策略1: 优先检查大纲结构和标题样式
            headings_from_styles = self._extract_headings_from_word_styles(doc)
            
            # 策略2: 如果样式解析效果不好，尝试文本分析
            headings_from_text = []
            if not headings_from_styles:
                print("样式解析未找到标题，尝试文本分析...")
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                headings_from_text = self._extract_headings_from_text(full_text)
            
            # 选择最佳策略
            if headings_from_styles:
                result['headings'] = headings_from_styles
                result['extraction_method'] = 'style_based'
                print(f"使用样式解析，找到{len(headings_from_styles)}个标题")
            elif headings_from_text:
                result['headings'] = headings_from_text
                result['extraction_method'] = 'text_analysis'
                print(f"使用文本分析，找到{len(headings_from_text)}个标题")
            else:
                result['headings'] = []
                result['extraction_method'] = 'no_headings_found'
                print("未找到任何标题结构")
            
            # 生成内容预览
            content_preview = ""
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and self._get_heading_level(para) == 0:  # 只取非标题内容
                    content_preview += text + " "
                    if len(content_preview) > 500:
                        break
            
            result['structure'] = self._build_document_structure(result['headings'])
            result['content_preview'] = content_preview[:500] + "..." if len(content_preview) > 500 else content_preview
            
            print(f"=== DOCX解析完成，最终找到{len(result['headings'])}个标题 ===\n")
            
        except Exception as e:
            print(f"DOCX解析错误: {str(e)}")
            raise Exception(f"DOCX解析错误: {str(e)}")
        
        return result
    
    def _extract_headings_from_word_styles(self, doc) -> List[Dict[str, Any]]:
        """从Word文档的样式和格式中提取标题"""
        headings = []
        
        print("检查Word文档的样式和格式...")
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # 获取标题级别
            heading_level = self._get_heading_level(para)
            
            if heading_level > 0:
                style_name = para.style.name if para.style else 'Unknown'
                headings.append({
                    'text': text,
                    'level': heading_level,
                    'style': style_name,
                    'paragraph_index': i
                })
                print(f"  找到标题: 级别{heading_level}, 样式'{style_name}', 内容: {text[:50]}")
        
        # 验证标题层级的合理性
        if headings:
            headings = self._validate_and_fix_heading_levels(headings)
        
        return headings
    
    def _validate_and_fix_heading_levels(self, headings: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """验证和修复标题层级的合理性"""
        if not headings:
            return headings
        
        print("验证标题层级合理性...")
        
        # 确保第一个标题从1级开始（如果第一个标题级别过高，调整所有标题）
        min_level = min(h['level'] for h in headings)
        if min_level > 1:
            level_offset = min_level - 1
            print(f"调整标题级别，所有标题级别减少{level_offset}")
            for heading in headings:
                heading['level'] = max(1, heading['level'] - level_offset)
        
        # 检查级别跳跃是否合理（不应该从1级直接跳到4级）
        prev_level = 0
        for i, heading in enumerate(headings):
            current_level = heading['level']
            
            # 如果级别跳跃过大（超过1级），调整当前级别
            if prev_level > 0 and current_level - prev_level > 1:
                adjusted_level = prev_level + 1
                print(f"标题级别跳跃过大: '{heading['text'][:30]}' 从{current_level}调整为{adjusted_level}")
                heading['level'] = adjusted_level
                current_level = adjusted_level
            
            prev_level = current_level
        
        return headings
    
    def _get_heading_level(self, paragraph) -> int:
        """获取段落的标题级别 - 优化版本"""
        text = paragraph.text.strip()
        if not text:
            return 0
        
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        # 1. 优先检查标准标题样式
        if 'heading' in style_name:
            try:
                # 提取标题级别数字
                match = re.search(r'heading\s*(\d+)', style_name)
                if match:
                    level = int(match.group(1))
                    print(f"发现标题样式: {style_name} -> 级别 {level}, 内容: {text[:50]}")
                    return min(level, 7)
                else:
                    print(f"标题样式但无级别: {style_name}, 内容: {text[:50]}")
                    return 1
            except Exception as e:
                print(f"解析标题样式失败: {style_name}, 错误: {e}")
                return 1
        
        # 2. 检查大纲级别（outline level）
        if hasattr(paragraph, '_element') and hasattr(paragraph._element, 'get_outline_level'):
            try:
                outline_level = paragraph._element.get_outline_level()
                if outline_level is not None and outline_level > 0:
                    print(f"发现大纲级别: {outline_level}, 内容: {text[:50]}")
                    return min(outline_level, 7)
            except:
                pass
        
        # 3. 检查段落格式的大纲级别
        if hasattr(paragraph, 'paragraph_format') and hasattr(paragraph.paragraph_format, 'outline_level'):
            try:
                outline_level = paragraph.paragraph_format.outline_level
                if outline_level is not None and outline_level > 0:
                    print(f"发现段落大纲级别: {outline_level}, 内容: {text[:50]}")
                    return min(outline_level, 7)
            except:
                pass
        
        # 4. 检查字体样式特征
        if paragraph.runs:
            for run in paragraph.runs:
                # 检查是否加粗
                is_bold = run.bold
                # 检查字体大小
                font_size = None
                if run.font.size:
                    font_size = run.font.size.pt
                
                # 基于字体特征判断
                if is_bold and font_size:
                    if font_size >= 18:
                        print(f"字体特征判断为1级标题: 大小{font_size}, 加粗, 内容: {text[:50]}")
                        return 1
                    elif font_size >= 16:
                        print(f"字体特征判断为2级标题: 大小{font_size}, 加粗, 内容: {text[:50]}")
                        return 2
                    elif font_size >= 14:
                        print(f"字体特征判断为3级标题: 大小{font_size}, 加粗, 内容: {text[:50]}")
                        return 3
                elif is_bold:
                    # 仅加粗，可能是4级标题
                    print(f"仅加粗判断为4级标题: 内容: {text[:50]}")
                    return 4
        
        # 5. 最后检查文本模式（更严格的匹配）
        # 只有在长度合适且符合标题特征时才使用文本模式
        if len(text) <= 150:  # 标题通常不会太长
            # 首先检查主要模式
            for i, pattern in enumerate(self.heading_patterns):
                if re.match(pattern, text):
                    level = min(i + 1, 7)
                    print(f"主要模式匹配第{i+1}个模式，判断为{level}级标题: {text[:50]}")
                    return level
            
            # 如果主要模式没匹配，检查次要模式（起始级别为5）
            for i, pattern in enumerate(self.secondary_patterns):
                if re.match(pattern, text):
                    level = min(i + 5, 7)  # 次要模式从5级开始
                    print(f"次要模式匹配第{i+1}个模式，判断为{level}级标题: {text[:50]}")
                    return level
        
        return 0
    
    def _extract_headings_from_bookmarks(self, pdf_reader) -> List[Dict[str, Any]]:
        """从PDF书签（outline）中提取标题结构"""
        headings = []
        
        try:
            # 获取PDF的outline/书签
            outline = pdf_reader.outline
            if not outline:
                return headings
            
            print(f"Outline type: {type(outline)}")
            print(f"Outline length: {len(outline) if hasattr(outline, '__len__') else 'N/A'}")
            
            def extract_bookmark_info(bookmark_item, level=1):
                """提取单个书签信息"""
                try:
                    title = None
                    
                    # 尝试获取标题
                    if hasattr(bookmark_item, 'title'):
                        title = bookmark_item.title
                    elif isinstance(bookmark_item, dict):
                        title = bookmark_item.get('/Title', bookmark_item.get('title'))
                    
                    if title and isinstance(title, str):
                        title = title.strip()
                        if title:
                            page_num = None
                            try:
                                page_num = self._get_bookmark_page(bookmark_item, pdf_reader)
                            except:
                                pass
                            
                            heading_data = {
                                'text': title,
                                'level': min(level, 7),
                                'style': f'Bookmark Level {min(level, 7)}'
                            }
                            
                            if page_num:
                                heading_data['page'] = page_num
                            
                            headings.append(heading_data)
                            print(f"Level {level}: {title}")
                            
                except Exception as e:
                    print(f"提取书签信息出错: {str(e)}")
            
            def traverse_outline(items, level=1):
                """遍历outline结构"""
                if not items:
                    return
                
                try:
                    for item in items:
                        if isinstance(item, list):
                            # 如果是列表，递归处理，层级增加
                            traverse_outline(item, level + 1)
                        else:
                            # 处理书签对象
                            extract_bookmark_info(item, level)
                            
                except Exception as e:
                    print(f"遍历outline出错: {str(e)}")
            
            # 开始处理
            traverse_outline(outline, 1)
            
        except Exception as e:
            print(f"书签解析失败: {str(e)}")
            import traceback
            traceback.print_exc()
            return []
        
        return headings
    
    def _get_bookmark_page(self, bookmark, pdf_reader):
        """获取书签对应的页码"""
        try:
            if hasattr(bookmark, 'page'):
                if hasattr(bookmark.page, 'idnum'):
                    # 查找页码
                    for page_num, page in enumerate(pdf_reader.pages):
                        if hasattr(page, 'indirect_ref') and hasattr(page.indirect_ref, 'idnum'):
                            if page.indirect_ref.idnum == bookmark.page.idnum:
                                return page_num + 1
            return None
        except:
            return None
    
    def _extract_pdf_bookmarks_intelligent(self, pdf_reader) -> List[Dict[str, Any]]:
        """智能PDF书签解析方法"""
        headings = []
        
        try:
            outline = pdf_reader.outline
            if not outline:
                print("PDF文档无书签结构")
                return headings
            
            print("发现PDF书签，开始智能解析...")
            
            # 尝试多种解析方法
            methods = [
                self._extract_bookmarks_method_1,
                self._extract_bookmarks_method_2,
                self._extract_headings_from_bookmarks_simple
            ]
            
            for i, method in enumerate(methods, 1):
                try:
                    print(f"尝试书签解析方法 {i}...")
                    result = method(pdf_reader)
                    if result and len(result) > 0:
                        print(f"方法 {i} 成功，找到 {len(result)} 个标题")
                        # 验证结果质量
                        if self._validate_bookmark_quality(result):
                            return self._optimize_bookmark_levels(result)
                        else:
                            print(f"方法 {i} 结果质量不佳，尝试下一种方法")
                            continue
                except Exception as e:
                    print(f"方法 {i} 失败: {str(e)}")
                    continue
            
            print("所有书签解析方法都失败")
            return []
            
        except Exception as e:
            print(f"书签解析整体失败: {str(e)}")
            return []
    
    def _extract_bookmarks_method_1(self, pdf_reader) -> List[Dict[str, Any]]:
        """书签解析方法1：标准递归解析"""
        headings = []
        
        def process_outline_recursive(items, level=1):
            if not items:
                return
            
            for item in items:
                if isinstance(item, list):
                    process_outline_recursive(item, level + 1)
                else:
                    try:
                        title = None
                        if hasattr(item, 'title') and item.title:
                            title = str(item.title).strip()
                        
                        if title and len(title) <= 200:  # 合理的标题长度
                            headings.append({
                                'text': title,
                                'level': min(level, 7),
                                'style': f'Bookmark Level {level}',
                                'method': 'recursive'
                            })
                    except:
                        continue
        
        process_outline_recursive(pdf_reader.outline)
        return headings
    
    def _extract_bookmarks_method_2(self, pdf_reader) -> List[Dict[str, Any]]:
        """书签解析方法2：扁平化解析"""
        headings = []
        
        def flatten_all_items(items, base_level=1):
            for item in items:
                try:
                    if hasattr(item, 'title') and item.title:
                        title = str(item.title).strip()
                        if title and len(title) <= 200:
                            headings.append({
                                'text': title,
                                'level': base_level,
                                'style': f'Bookmark Level {base_level}',
                                'method': 'flatten'
                            })
                            base_level += 1
                    
                    if isinstance(item, list):
                        flatten_all_items(item, base_level)
                except:
                    continue
        
        flatten_all_items(pdf_reader.outline)
        return headings
    
    def _extract_headings_from_bookmarks_simple(self, pdf_reader) -> List[Dict[str, Any]]:
        """简化版书签解析方法（保持原有逻辑）"""
        headings = []
        
        try:
            outline = pdf_reader.outline
            if not outline:
                return headings
            
            def flatten_outline(items, level=1):
                for item in items:
                    try:
                        if hasattr(item, 'title') and item.title:
                            title = str(item.title).strip()
                            if title:
                                headings.append({
                                    'text': title,
                                    'level': min(level, 7),
                                    'style': f'Bookmark Level {min(level, 7)}',
                                    'method': 'simple'
                                })
                        
                        if isinstance(item, list):
                            flatten_outline(item, level + 1)
                    except:
                        continue
            
            flatten_outline(outline)
            
        except Exception as e:
            print(f"简化书签解析失败: {str(e)}")
            return []
        
        return headings
    
    def _validate_bookmark_quality(self, headings: List[Dict[str, Any]]) -> bool:
        """验证书签解析结果的质量"""
        if not headings:
            return False
        
        # 检查是否有合理的标题数量
        if len(headings) < 2:
            print("书签数量过少，可能不是有效的文档结构")
            return False
        
        # 检查是否有过长的"标题"（可能是内容片段）
        long_titles = [h for h in headings if len(h['text']) > 100]
        if len(long_titles) > len(headings) * 0.5:
            print("标题过长的比例过高，可能解析错误")
            return False
        
        # 检查是否有层级结构
        levels = [h['level'] for h in headings]
        if len(set(levels)) == 1 and len(headings) > 10:
            print("所有标题都是同一级别且数量很多，可能解析错误")
            return False
        
        print("书签质量验证通过")
        return True
    
    def _optimize_bookmark_levels(self, headings: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """优化书签层级"""
        if not headings:
            return headings
        
        print("优化书签层级...")
        
        # 如果所有书签都是同一级别，尝试通过内容推断层级
        levels = [h['level'] for h in headings]
        if len(set(levels)) == 1:
            print("所有书签同级别，尝试通过内容推断层级")
            for i, heading in enumerate(headings):
                text = heading['text']
                # 根据文本特征调整级别
                if any(pattern in text for pattern in ['第一章', '第二章', '第三章', '第四章', '第五章']):
                    heading['level'] = 1
                elif any(pattern in text for pattern in ['第一节', '第二节', '第三节', '第四节', '第五节']):
                    heading['level'] = 2
                elif text.startswith(('1.', '2.', '3.', '4.', '5.')):
                    heading['level'] = 2
                elif text.startswith(('1.1', '1.2', '2.1', '2.2', '3.1')):
                    heading['level'] = 3
        
        return headings

    def _extract_headings_from_text(self, text: str) -> List[Dict[str, Any]]:
        """从纯文本中提取标题 - 改进版本"""
        headings = []
        lines = text.split('\n')
        
        print("开始文本模式标题提取...")
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line or len(line) > 200:  # 跳过空行和过长的行
                continue
            
            # 首先检查主要模式
            matched = False
            for level, pattern in enumerate(self.heading_patterns, 1):
                if re.match(pattern, line):
                    headings.append({
                        'text': line,
                        'level': min(level, 7),
                        'style': f'Text Pattern L{min(level, 7)}',
                        'line_number': line_num + 1
                    })
                    print(f"  文本标题: L{min(level, 7)} - {line[:50]}")
                    matched = True
                    break
            
            # 如果主要模式没匹配，检查次要模式
            if not matched:
                for level, pattern in enumerate(self.secondary_patterns, 5):
                    if re.match(pattern, line):
                        headings.append({
                            'text': line,
                            'level': min(level, 7),
                            'style': f'Text Pattern L{min(level, 7)}',
                            'line_number': line_num + 1
                        })
                        print(f"  次要标题: L{min(level, 7)} - {line[:50]}")
                        break
        
        # 验证和优化文本提取的结果
        if headings:
            headings = self._validate_and_fix_text_headings(headings)
        
        print(f"文本模式提取完成，找到{len(headings)}个标题")
        return headings
    
    def _validate_and_fix_text_headings(self, headings: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """验证和修复文本提取的标题"""
        if not headings:
            return headings
        
        print("验证文本提取的标题...")
        
        # 移除重复的标题
        seen_texts = set()
        unique_headings = []
        for heading in headings:
            text = heading['text'].strip()
            if text not in seen_texts:
                seen_texts.add(text)
                unique_headings.append(heading)
            else:
                print(f"移除重复标题: {text[:30]}")
        
        # 检查标题质量
        valid_headings = []
        for heading in unique_headings:
            text = heading['text']
            
            # 过滤掉可能的误匹配
            if self._is_likely_heading(text):
                valid_headings.append(heading)
            else:
                print(f"过滤无效标题: {text[:30]}")
        
        return valid_headings
    
    def _is_likely_heading(self, text: str) -> bool:
        """判断文本是否可能是标题"""
        # 长度检查
        if len(text) < 3 or len(text) > 150:
            return False
        
        # 检查是否包含过多数字（可能是数据）
        digit_count = sum(1 for c in text if c.isdigit())
        if digit_count > len(text) * 0.5:
            return False
        
        # 检查是否包含常见的非标题特征
        non_heading_patterns = [
            r'^\d+$',  # 纯数字
            r'^[.。,，;；]+$',  # 纯标点
            r'[：:]$',  # 以冒号结尾（可能是列表项）
            r'^(注|备注|说明)[：:]',  # 注释文本
            r'[。.!！]$',  # 以句号感叹号结尾（可能是正文）
        ]
        
        for pattern in non_heading_patterns:
            if re.search(pattern, text):
                return False
        
        return True
    
    def _build_document_structure(self, headings: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """构建文档结构树"""
        if not headings:
            return []
        
        structure = []
        stack = []
        
        for heading in headings:
            level = heading['level']
            
            # 清理栈，保持层级关系
            while stack and stack[-1]['level'] >= level:
                stack.pop()
            
            # 创建节点
            node = {
                'text': heading['text'],
                'level': level,
                'style': heading.get('style', ''),
                'children': []
            }
            
            # 添加到父节点或根节点
            if stack:
                stack[-1]['children'].append(node)
            else:
                structure.append(node)
            
            stack.append(node)
        
        return structure
    
    def get_statistics(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """获取文档统计信息"""
        headings = result.get('headings', [])
        
        stats = {
            'total_headings': len(headings),
            'heading_levels': {},
            'max_level': 0,
            'document_type': result.get('document_type', 'Unknown')
        }
        
        for heading in headings:
            level = heading['level']
            stats['heading_levels'][level] = stats['heading_levels'].get(level, 0) + 1
            stats['max_level'] = max(stats['max_level'], level)
        
        if result['document_type'] == 'PDF':
            stats['total_pages'] = result.get('total_pages', 0)
        else:
            stats['total_paragraphs'] = result.get('total_paragraphs', 0)
        
        return stats 