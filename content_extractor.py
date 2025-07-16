import os
import re
from typing import Dict, List, Any, Optional
from document_parser import DocumentParser
import PyPDF2
from docx import Document

class ContentExtractor:
    """内容提取器 - 根据标题和关键词提取文档内容"""
    
    def __init__(self):
        self.parser = DocumentParser()
        
    def extract_content_by_title_and_keywords(self, document_path: str, 
                                            document_structure: Dict[str, Any],
                                            target_title: str,
                                            keywords: List[str]) -> Optional[Dict[str, Any]]:
        """
        根据目标标题和关键词提取文档内容
        
        Args:
            document_path: 文档路径
            document_structure: 文档结构信息
            target_title: 目标标题
            keywords: 关键词列表
            
        Returns:
            提取的内容信息，包含content, start_heading, end_heading, confidence
        """
        print(f"\n--- 提取目标: {target_title} ---")
        print(f"关键词: {', '.join(keywords)}")
        
        file_extension = os.path.splitext(document_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(document_path, document_structure, target_title, keywords)
        elif file_extension in ['.doc', '.docx']:
            return self._extract_from_docx(document_path, document_structure, target_title, keywords)
        else:
            print(f"不支持的文件格式: {file_extension}")
            return None
    
    def _extract_from_pdf(self, document_path: str, document_structure: Dict[str, Any],
                         target_title: str, keywords: List[str]) -> Optional[Dict[str, Any]]:
        """从PDF文档中提取内容"""
        try:
            with open(document_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # 提取所有文本内容
                full_text = ""
                page_texts = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        page_texts.append({
                            'page_num': page_num + 1,
                            'text': page_text
                        })
                        full_text += page_text + "\n"
                    except Exception as e:
                        print(f"提取第{page_num + 1}页失败: {e}")
                        continue
                
                # 基于标题结构和关键词提取内容
                return self._extract_content_by_structure_and_keywords(
                    full_text, page_texts, document_structure, target_title, keywords
                )
                
        except Exception as e:
            print(f"PDF内容提取失败: {e}")
            return None
    
    def _extract_from_docx(self, document_path: str, document_structure: Dict[str, Any],
                          target_title: str, keywords: List[str]) -> Optional[Dict[str, Any]]:
        """从DOCX文档中提取内容"""
        try:
            doc = Document(document_path)
            
            # 提取所有段落文本
            full_text = ""
            paragraphs_info = []
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs_info.append({
                        'index': i,
                        'text': text,
                        'style': para.style.name if para.style else 'Normal'
                    })
                    full_text += text + "\n"
            
            # 基于标题结构和关键词提取内容
            return self._extract_content_by_structure_and_keywords(
                full_text, paragraphs_info, document_structure, target_title, keywords
            )
            
        except Exception as e:
            print(f"DOCX内容提取失败: {e}")
            return None
    
    def _extract_content_by_structure_and_keywords(self, full_text: str, 
                                                 text_parts: List[Dict],
                                                 document_structure: Dict[str, Any],
                                                 target_title: str,
                                                 keywords: List[str]) -> Optional[Dict[str, Any]]:
        """
        基于文档结构和关键词提取内容
        
        Args:
            full_text: 完整文本
            text_parts: 文本部分列表（页面或段落）
            document_structure: 文档结构
            target_title: 目标标题
            keywords: 关键词列表
            
        Returns:
            提取的内容信息
        """
        headings = document_structure.get('headings', [])
        
        # 策略1: 精确匹配标题
        exact_match_result = self._find_content_by_exact_title_match(
            full_text, headings, target_title, keywords
        )
        if exact_match_result:
            print(f"✓ 通过精确标题匹配找到内容")
            return exact_match_result
        
        # 策略2: 模糊匹配标题
        fuzzy_match_result = self._find_content_by_fuzzy_title_match(
            full_text, headings, target_title, keywords
        )
        if fuzzy_match_result:
            print(f"✓ 通过模糊标题匹配找到内容")
            return fuzzy_match_result
        
        # 策略3: 关键词匹配
        keyword_match_result = self._find_content_by_keywords(
            full_text, headings, keywords
        )
        if keyword_match_result:
            print(f"✓ 通过关键词匹配找到内容")
            return keyword_match_result
        
        # 策略4: 智能语义匹配
        semantic_match_result = self._find_content_by_semantic_match(
            full_text, headings, target_title, keywords
        )
        if semantic_match_result:
            print(f"✓ 通过语义匹配找到内容")
            return semantic_match_result
        
        print(f"✗ 未找到匹配的内容")
        return None
    
    def _find_content_by_exact_title_match(self, full_text: str, headings: List[Dict],
                                         target_title: str, keywords: List[str]) -> Optional[Dict[str, Any]]:
        """通过精确标题匹配查找内容"""
        print(f"尝试精确匹配标题: '{target_title}'")
        
        # 首先尝试完全匹配
        for i, heading in enumerate(headings):
            heading_text = heading['text']
            
            # 完全匹配（去除空格和标点符号的影响）
            title_clean = re.sub(r'[\s\u3000]+', '', target_title.lower())
            heading_clean = re.sub(r'[\s\u3000]+', '', heading_text.lower())
            
            if title_clean == heading_clean:
                print(f"✓ 找到完全匹配的标题: '{heading_text}'")
                return self._extract_content_between_headings(
                    full_text, headings, i, heading_text
                )
        
        # 其次尝试包含匹配（target_title包含在heading_text中，或vice versa）
        for i, heading in enumerate(headings):
            heading_text = heading['text']
            
            # 去除数字编号前缀，提取核心标题内容
            target_core = re.sub(r'^[\d\.\s]+', '', target_title).strip()
            heading_core = re.sub(r'^[\d\.\s]+', '', heading_text).strip()
            
            if len(target_core) >= 3 and len(heading_core) >= 3:
                if target_core.lower() in heading_core.lower() or heading_core.lower() in target_core.lower():
                    print(f"✓ 找到包含匹配的标题: '{heading_text}' (目标核心: '{target_core}')")
                    return self._extract_content_between_headings(
                        full_text, headings, i, heading_text
                    )
        
        print(f"✗ 未找到精确匹配的标题")
        return None
    
    def _find_content_by_fuzzy_title_match(self, full_text: str, headings: List[Dict],
                                         target_title: str, keywords: List[str]) -> Optional[Dict[str, Any]]:
        """通过模糊标题匹配查找内容"""
        print(f"尝试模糊匹配标题: '{target_title}'")
        best_match_score = 0
        best_match_index = -1
        best_heading_text = ""
        
        for i, heading in enumerate(headings):
            heading_text = heading['text']
            
            # 计算匹配分数
            score = self._calculate_match_score(heading_text, target_title, keywords)
            
            # 降低阈值，确保能找到相关内容，但优先选择标题匹配分数高的
            if score > best_match_score and score > 0.2:  # 降低阈值从0.3到0.2
                best_match_score = score
                best_match_index = i
                best_heading_text = heading_text
        
        if best_match_index >= 0:
            print(f"✓ 找到最佳匹配标题: '{best_heading_text}' (分数: {best_match_score:.3f})")
            return self._extract_content_between_headings(
                full_text, headings, best_match_index, best_heading_text
            )
        
        print(f"✗ 模糊匹配未找到合适的标题")
        return None
    
    def _find_content_by_keywords(self, full_text: str, headings: List[Dict],
                                keywords: List[str]) -> Optional[Dict[str, Any]]:
        """通过关键词匹配查找内容"""
        print(f"尝试关键词匹配: {keywords}")
        
        # 策略1: 在标题中查找包含关键词的章节
        best_heading_score = 0
        best_heading_index = -1
        best_heading_text = ""
        
        for i, heading in enumerate(headings):
            heading_text = heading['text']
            heading_lower = heading_text.lower()
            
            # 计算该标题包含多少个关键词
            keyword_count = sum(1 for keyword in keywords if keyword.lower() in heading_lower)
            
            if keyword_count > best_heading_score:
                best_heading_score = keyword_count
                best_heading_index = i
                best_heading_text = heading_text
        
        # 如果找到包含关键词的标题，提取该标题下的内容
        if best_heading_score >= 1:  # 至少包含1个关键词
            print(f"✓ 找到包含{best_heading_score}个关键词的标题: '{best_heading_text}'")
            return self._extract_content_between_headings(
                full_text, headings, best_heading_index, best_heading_text
            )
        
        # 策略2: 如果标题中没有找到，在文档内容中查找关键词密集区域
        print("在标题中未找到关键词，搜索内容中的关键词密集区域...")
        lines = full_text.split('\n')
        best_region_start = -1
        best_region_score = 0
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            score = sum(1 for keyword in keywords if keyword.lower() in line_lower)
            
            if score > best_region_score and score >= 2:  # 至少包含2个关键词
                best_region_score = score
                best_region_start = i
        
        if best_region_start >= 0:
            # 提取该区域周围的内容（前后各20行）
            start_line = max(0, best_region_start - 20)
            end_line = min(len(lines), best_region_start + 50)
            
            content = '\n'.join(lines[start_line:end_line])
            
            print(f"✓ 找到关键词密集区域（第{best_region_start + 1}行附近，包含{best_region_score}个关键词）")
            return {
                'content': content,
                'start_heading': f"关键词匹配区域（第{best_region_start + 1}行附近）",
                'end_heading': f"区域结束（第{end_line}行）",
                'confidence': 0.6 + (best_region_score * 0.1)
            }
        
        print(f"✗ 关键词匹配未找到相关内容")
        return None
    
    def _find_content_by_semantic_match(self, full_text: str, headings: List[Dict],
                                      target_title: str, keywords: List[str]) -> Optional[Dict[str, Any]]:
        """通过语义匹配查找内容（基于规则的简单实现）"""
        # 创建语义匹配词典
        semantic_groups = {
            '项目信息': ['项目', '工程', '名称', '编号', '概况', '简介'],
            '资格要求': ['资格', '资质', '条件', '要求', '标准', '规定'],
            '投标须知': ['须知', '注意', '事项', '说明', '规定', '要求'],
            '保证金': ['保证金', '担保', '金额', '缴纳', '保证'],
            '商务条件': ['商务', '合同', '条款', '付款', '价格'],
            '技术要求': ['技术', '规格', '参数', '功能', '性能'],
            '废标条件': ['废标', '无效', '拒绝', '不符合']
        }
        
        # 找到最匹配的语义组
        target_group = None
        for group_name, group_keywords in semantic_groups.items():
            if any(keyword in target_title for keyword in group_keywords):
                target_group = group_keywords
                break
        
        if target_group:
            # 扩展关键词列表
            extended_keywords = keywords + target_group
            
            # 基于扩展关键词重新搜索
            return self._find_content_by_keywords(full_text, headings, extended_keywords)
        
        return None
    
    def _extract_content_between_headings(self, full_text: str, headings: List[Dict],
                                        heading_index: int, heading_text: str) -> Dict[str, Any]:
        """提取指定标题到下一个同级或更高级标题之间的内容"""
        print(f"开始提取标题内容: '{heading_text}'")
        
        current_heading = headings[heading_index]
        current_level = current_heading['level']
        
        # 找到下一个同级或更高级的标题
        next_heading_text = ""
        next_heading_index = -1
        
        for i in range(heading_index + 1, len(headings)):
            if headings[i]['level'] <= current_level:
                next_heading_text = headings[i]['text']
                next_heading_index = i
                break
        
        print(f"下一个标题: '{next_heading_text}' (索引: {next_heading_index})")
        
        # 策略1: 尝试多种标题匹配方式
        heading_match = None
        start_pos = -1
        
        # 方式1: 精确匹配（转义特殊字符）
        try:
            heading_pattern = re.escape(heading_text.strip())
            heading_match = re.search(heading_pattern, full_text, re.IGNORECASE)
            if heading_match:
                start_pos = heading_match.end()
                print(f"✓ 精确匹配成功，起始位置: {start_pos}")
        except:
            pass
        
        # 方式2: 部分匹配（去除特殊字符）
        if not heading_match:
            # 创建更宽松的匹配模式
            clean_heading = re.sub(r'[（）\(\)\[\]【】《》""''`~!@#$%^&*+=|\\:;\"\'<>,.?/]', '', heading_text).strip()
            words = clean_heading.split()
            if words:
                # 使用第一个有意义的词
                for word in words:
                    if len(word) >= 2:  # 至少2个字符
                        pattern = re.escape(word)
                        matches = list(re.finditer(pattern, full_text, re.IGNORECASE))
                        for match in matches:
                            # 检查匹配位置前后的上下文是否像标题
                            context_start = max(0, match.start() - 20)
                            context_end = min(len(full_text), match.end() + 20)
                            context = full_text[context_start:context_end]
                            
                            # 简单判断：如果前面有换行符，可能是标题
                            if '\n' in full_text[context_start:match.start()]:
                                heading_match = match
                                start_pos = match.start()
                                print(f"✓ 部分匹配成功（词汇: '{word}'），起始位置: {start_pos}")
                                break
                        if heading_match:
                            break
        
        # 方式3: 关键词组合匹配
        if not heading_match:
            # 提取标题中的关键词
            keywords = re.findall(r'[\u4e00-\u9fff]{2,}', heading_text)  # 提取中文词汇
            if keywords:
                # 查找包含多个关键词的行
                lines = full_text.split('\n')
                for i, line in enumerate(lines):
                    line_lower = line.lower()
                    heading_lower = heading_text.lower()
                    
                    # 计算关键词匹配数
                    keyword_matches = sum(1 for kw in keywords if kw.lower() in line_lower)
                    
                    # 如果匹配了足够多的关键词，认为找到了标题
                    if keyword_matches >= min(2, len(keywords)):  # 至少匹配2个关键词或全部关键词
                        # 计算该行在full_text中的位置
                        line_start = full_text.find(line)
                        if line_start >= 0:
                            start_pos = line_start + len(line)
                            print(f"✓ 关键词匹配成功（匹配{keyword_matches}个关键词），起始位置: {start_pos}")
                            break
        
        if start_pos < 0:
            print(f"✗ 未找到标题 '{heading_text}' 在文档中的位置")
            return None
        
        # 找到结束位置
        end_pos = len(full_text)  # 默认到文档结尾
        
        if next_heading_text:
            # 寻找下一个标题的位置
            next_match = None
            
            # 方式1: 精确匹配下一个标题
            try:
                next_pattern = re.escape(next_heading_text.strip())
                next_match = re.search(next_pattern, full_text[start_pos:], re.IGNORECASE)
            except:
                pass
            
            # 方式2: 部分匹配下一个标题
            if not next_match:
                next_clean = re.sub(r'[（）\(\)\[\]【】《》""''`~!@#$%^&*+=|\\:;\"\'<>,.?/]', '', next_heading_text).strip()
                next_words = next_clean.split()
                if next_words:
                    for word in next_words:
                        if len(word) >= 2:
                            pattern = re.escape(word)
                            next_match = re.search(pattern, full_text[start_pos:], re.IGNORECASE)
                            if next_match:
                                break
            
            if next_match:
                end_pos = start_pos + next_match.start()
                print(f"✓ 找到结束位置: {end_pos} (下一个标题: '{next_heading_text}')")
        
        # 提取内容
        content = full_text[start_pos:end_pos].strip()
        
        if not content:
            print(f"✗ 提取的内容为空")
            return None
        
        # 清理内容
        content = self._clean_extracted_content(content)
        
        print(f"✓ 成功提取内容: {len(content)} 字符")
        print(f"内容预览: {content[:100]}...")
        
        return {
            'content': content,
            'start_heading': heading_text,
            'end_heading': next_heading_text or "文档结尾",
            'confidence': 0.9
        }
    
    def _calculate_match_score(self, heading_text: str, target_title: str, keywords: List[str]) -> float:
        """计算标题匹配分数"""
        heading_lower = heading_text.lower()
        target_lower = target_title.lower()
        
        score = 0.0
        
        # 1. 标题直接匹配度（权重最高）
        # 移除编号和标点，比较核心内容
        target_core = re.sub(r'^[\d\.\s\u3000]+', '', target_lower).strip()
        heading_core = re.sub(r'^[\d\.\s\u3000]+', '', heading_lower).strip()
        
        if target_core and heading_core:
            # 完全匹配
            if target_core == heading_core:
                score += 1.0
            # 包含匹配
            elif target_core in heading_core or heading_core in target_core:
                score += 0.8
            else:
                # 词汇重叠度
                target_words = set(target_core.split())
                heading_words = set(heading_core.split())
                if target_words and heading_words:
                    overlap = len(target_words & heading_words)
                    total = len(target_words | heading_words)
                    score += (overlap / total) * 0.6
        
        # 2. 关键词匹配度（权重较低，仅作为辅助）
        if keywords:
            keyword_matches = sum(1 for keyword in keywords if keyword.lower() in heading_lower)
            score += (keyword_matches / len(keywords)) * 0.2
        
        print(f"标题匹配分析: '{heading_text}' vs '{target_title}' = {score:.3f}")
        return score
    
    def _clean_extracted_content(self, content: str) -> str:
        """清理提取的内容"""
        if not content:
            return ""
        
        # 移除多余的空白行
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # 只保留非空行
                cleaned_lines.append(line)
        
        # 限制内容长度（避免过长）
        cleaned_content = '\n'.join(cleaned_lines)
        if len(cleaned_content) > 5000:  # 限制5000字符
            cleaned_content = cleaned_content[:5000] + "\n...[内容过长，已截断]"
        
        return cleaned_content
    
    def extract_all_content_by_headings(self, document_path: str, 
                                      document_structure: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        提取文档中所有标题下的内容
        
        Args:
            document_path: 文档路径
            document_structure: 文档结构
            
        Returns:
            所有标题内容列表
        """
        print(f"\n=== 提取所有标题内容 ===")
        
        headings = document_structure.get('headings', [])
        if not headings:
            print("文档中没有找到标题")
            return []
        
        file_extension = os.path.splitext(document_path)[1].lower()
        
        # 获取完整文本
        if file_extension == '.pdf':
            full_text = self._get_pdf_full_text(document_path)
        elif file_extension in ['.doc', '.docx']:
            full_text = self._get_docx_full_text(document_path)
        else:
            print(f"不支持的文件格式: {file_extension}")
            return []
        
        if not full_text:
            return []
        
        all_contents = []
        
        for i, heading in enumerate(headings):
            content_info = self._extract_content_between_headings(
                full_text, headings, i, heading['text']
            )
            
            if content_info and content_info['content']:
                all_contents.append({
                    'heading': heading,
                    'content_info': content_info
                })
                print(f"✓ 提取标题: {heading['text']} ({len(content_info['content'])} 字符)")
            else:
                print(f"✗ 跳过标题: {heading['text']} (无内容)")
        
        print(f"总共提取了 {len(all_contents)} 个标题的内容")
        return all_contents
    
    def _get_pdf_full_text(self, document_path: str) -> str:
        """获取PDF完整文本"""
        try:
            with open(document_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                
                for page in pdf_reader.pages:
                    try:
                        full_text += page.extract_text() + "\n"
                    except:
                        continue
                
                return full_text
        except Exception as e:
            print(f"获取PDF文本失败: {e}")
            return ""
    
    def _get_docx_full_text(self, document_path: str) -> str:
        """获取DOCX完整文本"""
        try:
            doc = Document(document_path)
            full_text = ""
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text += text + "\n"
            
            return full_text
        except Exception as e:
            print(f"获取DOCX文本失败: {e}")
            return "" 